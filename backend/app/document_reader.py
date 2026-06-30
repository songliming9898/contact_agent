"""
文档读取模块 - Word(.docx) / PDF 文本+表格提取
"""
import re
from pathlib import Path
from typing import Dict, List, Tuple, Any


def read_docx(file_path: str) -> Dict[str, Any]:
    """
    读取 Word (.docx) 文件，提取文本和表格。
    
    Args:
        file_path: docx 文件路径
    
    Returns:
        {
            "full_text": str,       # 全文（保留段落结构）
            "tables": list,         # 表格列表 [二维数组]
            "table_contexts": list, # 表格上下文（表格前面的段落文本）
            "paragraphs": list,     # 段落文本列表
            "page_count": int,      # 估算页数
        }
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")
    
    doc = Document(file_path)
    
    paragraphs = []
    tables = []
    table_contexts = []
    
    for element in doc.element.body:
        # 获取元素类型
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        
        if tag == 'p':
            # 段落
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
        
        elif tag == 'tbl':
            # 表格
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break
            if table:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                if rows:
                    tables.append(rows)
                    # 取表格前最后一段作为上下文
                    context = paragraphs[-1] if paragraphs else ""
                    table_contexts.append(context)
    
    full_text = "\n".join(paragraphs)
    
    # 估算页数（约1500字符/页）
    page_count = max(1, len(full_text) // 1500)
    
    return {
        "full_text": full_text,
        "tables": tables,
        "table_contexts": table_contexts,
        "paragraphs": paragraphs,
        "page_count": page_count,
    }


def read_pdf(file_path: str) -> Dict[str, Any]:
    """
    读取 PDF 文件，提取文本和表格。
    
    Args:
        file_path: pdf 文件路径
    
    Returns:
        {
            "full_text": str,
            "tables": list,
            "table_contexts": list,
            "paragraphs": list,
            "page_count": int,
        }
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("请安装 PyMuPDF: pip install PyMuPDF")
    
    doc = fitz.open(file_path)
    
    all_text = []
    tables = []
    table_contexts = []
    
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            all_text.append(text)
        
        # 尝试检测表格（PyMuPDF 的表格检测）
        try:
            page_tables = page.find_tables()
            if page_tables and page_tables.tables:
                for tab in page_tables.tables:
                    rows = []
                    for row in tab.extract():
                        rows.append([str(cell).strip() if cell else "" for cell in row])
                    if rows and len(rows) > 1:
                        tables.append(rows)
                        # 取页面文本前部分作为上下文
                        context_lines = text.strip().split('\n')
                        context = context_lines[0] if context_lines else ""
                        table_contexts.append(context)
        except Exception:
            pass  # 表格检测失败，跳过
    
    doc.close()
    
    full_text = "\n".join(all_text)
    paragraphs = [p for p in full_text.split('\n') if p.strip()]
    page_count = len(doc)
    
    return {
        "full_text": full_text,
        "tables": tables,
        "table_contexts": table_contexts,
        "paragraphs": paragraphs,
        "page_count": page_count,
    }


def read_document(file_path: str) -> Dict[str, Any]:
    """
    统一文档读取入口，根据后缀分发。
    
    Args:
        file_path: 文档路径，支持 .docx / .pdf
    
    Returns:
        {
            "filename": str,
            "full_text": str,
            "tables": list,
            "table_contexts": list,
            "paragraphs": list,
            "page_count": int,
            "file_type": str,
        }
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    if suffix == '.docx':
        result = read_docx(file_path)
    elif suffix == '.pdf':
        result = read_pdf(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}，仅支持 .docx 和 .pdf")
    
    result["filename"] = path.name
    result["file_type"] = suffix.lstrip('.')
    
    # 基础清洗
    result["full_text"] = clean_text(result["full_text"])
    result["paragraphs"] = [clean_text(p) for p in result["paragraphs"] if clean_text(p).strip()]
    
    return result


def clean_text(text: str) -> str:
    """基础文本清洗：去除乱码、多余空白、不可见字符"""
    # 去除零宽字符
    text = re.sub(r'[\u200b\u200c\u200d\u200e\u200f\ufeff]', '', text)
    # 统一换行符
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # 去除连续空行（保留最多2个换行）
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 去除行首行尾空白
    text = '\n'.join(line.strip() for line in text.split('\n'))
    return text.strip()
